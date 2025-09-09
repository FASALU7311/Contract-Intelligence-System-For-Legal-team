
import os
os.environ["TRANSFORMERS_NO_TF"] = "1"

import streamlit as st
import torch
import pandas as pd
import fitz  # PyMuPDF
import re
from transformers import (
    BertTokenizer, BertForSequenceClassification
)
from streamlit_lottie import st_lottie
import requests
from streamlit_option_menu import option_menu
from fpdf import FPDF
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import random
from io import BytesIO
import docx
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import uuid
from datetime import datetime, timedelta
from streamlit_cookies_manager import EncryptedCookieManager
import time
import json
import os

st.set_page_config(page_title="Contract Intelligence System", layout="wide", page_icon="🗐")

# Get the directory of the script
script_dir = os.path.dirname(__file__)
css_path = os.path.join(script_dir, "style.css")

# Load CSS file
with open(css_path, "r", encoding="utf-8") as f:
    css = f.read()

# Inject CSS
st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

# ---------- Persistence helpers (disk + cookie) ----------
DATA_DIR = "persist"
os.makedirs(DATA_DIR, exist_ok=True)

def _analysis_path(key: str) -> str:
    return os.path.join(DATA_DIR, f"{key}.json")

def persist_analysis_to_disk(df: pd.DataFrame) -> str:
    user = st.session_state.get("user", "anon").replace("@", "_").replace(".", "_")
    key = f"{user}_{int(time.time())}_{uuid.uuid4().hex[:8]}"
    df.to_json(_analysis_path(key), orient="records")
    cookies["analysis_key"] = key
    cookies.save()
    st.session_state["analysis_key"] = key
    return key

def restore_analysis_from_cookie() -> pd.DataFrame | None:
    key = cookies.get("analysis_key") or st.session_state.get("analysis_key")
    if not key:
        return None
    path = _analysis_path(key)
    if os.path.exists(path):
        try:
            return pd.read_json(path)
        except Exception:
            return None
    return None

def clear_persisted_analysis():
    key = cookies.get("analysis_key") or st.session_state.get("analysis_key")
    if key:
        path = _analysis_path(key)
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass
    cookies["analysis_key"] = ""
    cookies.save()
    st.session_state.pop("analysis_key", None)
    st.session_state.pop("clause_data", None)

# --- Helper Functions ---
def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# --- Authentication UI ---
def register_page():
    st.subheader("Register")
    with st.form(key="register_form", clear_on_submit=True):
        username = st.text_input("Username", key="reg_username")
        email = st.text_input("Gmail Address", key="reg_email")
        password = st.text_input("Password", type="password", key="reg_pass")
        submit = st.form_submit_button("Send OTP")
        if submit:
            if not email.endswith("@gmail.com"):
                st.error("Please use a valid Gmail address.")
            elif user_exists(email):
                st.error("User already exists.")
            elif username_exists(username):
                st.error("Username already taken.")
            else:
                otp = generate_otp()
                if send_otp_email(email, otp):
                    df = load_users()
                    new_row = pd.DataFrame([{"username": username, "email": email, "password_hash": hash_password(password), "otp": otp, "is_verified": False}])
                    df = pd.concat([df, new_row], ignore_index=True)
                    save_users(df)
                    st.session_state["pending_email"] = email
                    st.success(f"OTP sent to {email}. Please verify.")
                else:
                    st.error("Failed to send OTP. Check email config.")

    if "pending_email" in st.session_state:
        with st.form(key="verify_form", clear_on_submit=True):
            otp_input = st.text_input("Enter OTP for Registration", key="reg_otp")
            if st.form_submit_button("Verify OTP"):
                user = get_user(st.session_state["pending_email"])
                if user is not None and str(user["otp"]) == otp_input:
                    verify_user(user["email"])
                    st.success("Registration complete! You can now login.")
                    del st.session_state["pending_email"]
                else:
                    st.error("Invalid OTP.")

# --- Cookie Manager ---
cookies = EncryptedCookieManager(
    prefix="contract_intel_",  
    password="super_secret_key_here"  # Change to secure random string
)
if not cookies.ready():
    st.stop()

# Default to upload page if no cookie exists
current_page = cookies.get("current_page", "Upload")

# --- Session State Restore ---
if "user" not in st.session_state:
    if cookies.get("user") and cookies.get("username"):
        st.session_state["user"] = cookies.get("user")
        st.session_state["username"] = cookies.get("username")
        st.session_state["authenticated"] = True
    else:
        st.session_state["authenticated"] = False

# --- LOGIN PAGE ---
def login_page():
    st.subheader("Login")
    with st.form(key="login_form", clear_on_submit=True):
        username = st.text_input("Username", key="login_username")
        password = st.text_input("Password", type="password", key="login_pass")
        col1, col2 = st.columns([2, 1])
        login_clicked = col1.form_submit_button("Login")
        forgot_clicked = col2.form_submit_button("Forgot Password")
        if login_clicked:
            user = get_user_by_username(username)
            if user is None:
                st.error("User not found.")
            elif user["password_hash"] != hash_password(password):
                st.error("Incorrect password.")
            elif not user["is_verified"]:
                st.error("User not verified. Please register and verify OTP.")
            else:
                st.session_state["authenticated"] = True
                st.session_state["user"] = user["email"]
                st.session_state["username"] = user["username"]
                cookies["user"] = user["email"]
                cookies["username"] = user["username"]
                cookies.save()
                st.success("Login successful!")
                time.sleep(0.5)
                st.rerun()
        if forgot_clicked:
            st.session_state["show_forgot_password"] = True
            st.session_state["authenticated"] = False
            st.rerun()

# --- LOGOUT ---
def logout():
    for key in ["authenticated", "user", "username", "show_forgot_password"]:
        if key in st.session_state:
            del st.session_state[key]
    cookies["user"] = ""
    cookies["username"] = ""
    cookies.save()
    st.success("Logged out successfully!")
    time.sleep(0.5)
    st.rerun()

def forgot_password_page():
    st.subheader("Forgot Password")
    if "fp_stage" not in st.session_state:
        st.session_state["fp_stage"] = "input_user_email"
    with st.form(key="forgot_form", clear_on_submit=True):
        if st.session_state["fp_stage"] == "input_user_email":
            username = st.text_input("Username", key="fp_username")
            email = st.text_input("Registered Gmail", key="fp_email")
            if st.form_submit_button("Send OTP for Reset"):
                user = get_user_by_username(username)
                if user is None:
                    st.error("User not found.")
                elif user["email"] != email:
                    st.error("Email does not match the username.")
                else:
                    otp = generate_otp()
                    update_user_otp(email, otp)
                    if send_otp_email(email, otp):
                        st.session_state["fp_pending_email"] = email
                        st.session_state["fp_stage"] = "otp_and_password"
                        st.success("OTP sent to your Gmail.")
                    else:
                        st.error("Failed to send OTP.")
        if st.session_state.get("fp_stage") == "otp_and_password":
            otp_input = st.text_input("Enter OTP for Password Reset", key="fp_otp")
            new_pass = st.text_input("New Password", type="password", key="fp_newpass")
            confirm_pass = st.text_input("Confirm New Password", type="password", key="fp_confirmpass")
            if st.form_submit_button("Reset Password"):
                user = get_user(st.session_state["fp_pending_email"])
                if user is not None and str(user["otp"]) == otp_input:
                    if new_pass != confirm_pass:
                        st.error("Passwords do not match.")
                    elif len(new_pass) < 6:
                        st.error("Password must be at least 6 characters.")
                    else:
                        update_user_password(user["email"], new_pass)
                        st.success("Password reset successful! You can now login.")
                        del st.session_state["fp_pending_email"]
                        st.session_state["fp_stage"] = "input_user_email"
                else:
                    st.error("Invalid OTP.")

# --- User Data Functions ---
USERS_CSV = "users.csv"

def load_users():
    cols = ["username", "email", "password_hash", "otp", "is_verified"]
    if not os.path.exists(USERS_CSV):
        return pd.DataFrame(columns=cols)
    df = pd.read_csv(USERS_CSV)
    for col in cols:
        if col not in df.columns:
            df[col] = "" if col != "is_verified" else False
    df = df[cols]
    return df

def save_users(df):
    df.to_csv(USERS_CSV, index=False)

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def generate_otp():
    return str(random.randint(100000, 999999))

def send_otp_email(receiver_email, otp):
    sender_email = "learnsaxon@gmail.com"  # <-- CHANGE THIS
    sender_password = "ihiumsydgpaceykb"
    subject = "Your OTP Code"
    body = f"Your OTP code is: {otp}"
    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))
    try:
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        print("Email send error:", e)
        return False

def user_exists(email):
    df = load_users()
    return email in df["email"].values

def username_exists(username):
    df = load_users()
    return username in df["username"].values

def get_user(email):
    df = load_users()
    user = df[df["email"] == email]
    return user.iloc[0] if not user.empty else None

def get_user_by_username(username):
    df = load_users()
    user = df[df["username"] == username]
    return user.iloc[0] if not user.empty else None

def update_user_otp(email, otp):
    df = load_users()
    df.loc[df["email"] == email, "otp"] = otp
    save_users(df)

def update_user_password(email, new_password):
    df = load_users()
    df.loc[df["email"] == email, "password_hash"] = hash_password(new_password)
    save_users(df)

def verify_user(email):
    df = load_users()
    df.loc[df["email"] == email, "is_verified"] = True
    save_users(df)

# --- Page Setup ---
lottie_contract = load_lottieurl("https://assets10.lottiefiles.com/packages/lf20_t9gkkhz4.json")

# --- Auth Navigation ---
if not st.session_state.get("authenticated", False):
    st.markdown('<div class="auth-card">', unsafe_allow_html=True)
    if st.session_state.get("show_forgot_password", False):
        forgot_password_page()
        if st.button("Back to Login"):
            st.session_state["show_forgot_password"] = False
            st.rerun()
    else:
        auth_tab = st.sidebar.radio("Authentication", ["Login", "Register", "Forgot Password"])
        if auth_tab == "Login":
            login_page()
        elif auth_tab == "Register":
            register_page()
        elif auth_tab == "Forgot Password":
            st.session_state["show_forgot_password"] = True
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()
else:
    st.sidebar.markdown(f"**Logged in as:** {st.session_state.get('username', '')}")
    if st.sidebar.button("Logout"):
        logout()

# --- Navigation Menu ---
if "current_page" not in st.session_state:
    st.session_state["current_page"] = cookies.get("current_page", "Upload")

with st.sidebar:
    selected = option_menu(
        "Navigation",
        ["Upload", "Analysis", "Insights", "Download"],
        icons=['cloud-upload', 'bar-chart', 'lightbulb', 'download'],
        menu_icon="cast",
        default_index=["Upload", "Analysis", "Insights", "Download"].index(st.session_state["current_page"]),
        key="main_navigation"
    )

if selected != st.session_state["current_page"]:
    st.session_state["current_page"] = selected
    cookies["current_page"] = selected
    cookies.save()
    st.rerun()
st.markdown('<h1>🗐 ℂ𝕠𝕟𝕥𝕣𝕒𝕔𝕥 𝕀𝕟𝕥𝕖𝕝𝕝𝕚𝕘𝕖𝕟𝕔𝕖 𝕊𝕪𝕤𝕥𝕖𝕞 𝕗𝕠𝕣 𝕃𝕖𝕘𝕒𝕝 𝕋𝕖𝕒𝕞𝕤</h1>', unsafe_allow_html=True) #####################################

# --- Load Legal-BERT Model ---
@st.cache_resource
def load_bert_model():
    model_path = "./legalbert_model"
    if not os.path.exists(model_path):
        st.error(f"Model directory '{model_path}' not found. Please ensure the model files are present.")
        return None, None
    try:
        bert_model = BertForSequenceClassification.from_pretrained(model_path)
        bert_tokenizer = BertTokenizer.from_pretrained(model_path)
        bert_model.eval()
        return bert_model, bert_tokenizer
    except Exception as e:
        st.error(f"Failed to load Legal-BERT model: {e}")
        return None, None

bert_model, bert_tokenizer = load_bert_model()
if not bert_model or not bert_tokenizer:
    st.stop()

# --- Load Legal Summarizer ---
@st.cache_resource
def load_pegasus_model():
    pegasus_model_name = "google/pegasus-xsum"
    try:
        pegasus_tokenizer = AutoTokenizer.from_pretrained(pegasus_model_name)
        pegasus_model = AutoModelForSeq2SeqLM.from_pretrained(pegasus_model_name)
        pegasus_model.eval()
        return pegasus_model, pegasus_tokenizer
    except Exception as e:
        st.error(f"Failed to load PEGASUS model: {e}")
        return None, None

pegasus_model, pegasus_tokenizer = load_pegasus_model()
if not pegasus_model or not pegasus_tokenizer:
    st.stop()

@st.cache_data
def load_label_mapping():
    if not os.path.exists("label_mapping.csv"):
        st.error("`label_mapping.csv` not found. Please provide the mapping file.")
        return None
    try:
        label_map_df = pd.read_csv("label_mapping.csv")
        label_map_df["clause_type"] = label_map_df["clause_type"].str.strip().str.replace('"', '')
        id2label = dict(label_map_df.values)
        return id2label
    except Exception as e:
        st.error(f"Error loading label mapping: {e}")
        return None

id2label = load_label_mapping()
if not id2label:
    st.stop()

# --- Clause Insight Mapping ---
clause_insights = {
    "Confidentiality": "Ensures sensitive information is protected.",
    "Termination for Convenience": "Allows parties to end the contract without cause.",
    "Dispute Resolution": "Specifies how disputes will be resolved.",
    "Cap on Liability": "Limits the amount a party has to pay if things go wrong.",
    "IP Ownership": "Clarifies who owns the intellectual property.",
    "Anti-Assignment": "Restricts transfer of contractual rights.",
    "Governing Law": "Specifies the legal jurisdiction governing the contract.",
    "Post-Termination": "Explains obligations even after the contract ends.",
    "General legal clause": "Standard legal clause without specific risk."
}

# --- Utils ---
@st.cache_data
def extract_deadlines_and_risks(text):
    text_lower = text.lower()
    deadline_patterns = [
        r"(?:within|at least|not less than|no later than)?\s*\(?\d{1,3}\)?\s*(?:days?|weeks?|months?|years?)'?\s*(?:prior to|before|after|preceding|from)?\s*(?:the\s+)?(?:effective\s+date|termination|claim|execution|delivery|notice|event)?",
        r"(?:on|by|before)?\s*(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+\d{1,2},?\s+\d{4}",
        r"(?:\d{1,2}/\d{1,2}/\d{2,4})"
    ]
    deadlines = []
    for pattern in deadline_patterns:
        found = re.findall(pattern, text, flags=re.IGNORECASE)
        for d in found:
            cleaned = d.strip().replace('\n', ' ')
            if len(cleaned) > 5 and not cleaned.lower().startswith(('or', 'and')):
                deadlines.append(cleaned)
    deadline_result = ", ".join(sorted(set(deadlines))) if deadlines else "-"
    risk_keywords = {
        "termination": "termination", "penalty": "penalty", "liability": "liability",
        "breach": "breach", "damages": "damages", "dispute": "dispute",
        "loss": "loss", "indemnity": "indemnity", "transfer": "transfer risk",
        "ownership": "IP risk"
    }
    risks = {label for word, label in risk_keywords.items() if word in text_lower}
    risk_result = ", ".join(sorted(risks)) if risks else "-"
    return deadline_result, risk_result

def generate_summaries_batch(clauses, batch_size=8):
    summaries = []
    for i in range(0, len(clauses), batch_size):
        batch_clauses = clauses[i:i+batch_size]
        filtered_batch = [clause if len(clause) > 100 else "[SHORT]" for clause in batch_clauses]
        inputs = pegasus_tokenizer(filtered_batch, truncation=True, padding="longest", return_tensors="pt")
        with torch.no_grad():
            summary_ids = pegasus_model.generate(**inputs, max_length=60, min_length=20, num_beams=4, early_stopping=True)
        decoded = pegasus_tokenizer.batch_decode(summary_ids, skip_special_tokens=True)
        for original, summary in zip(batch_clauses, decoded):
            summaries.append(original.strip() if len(original) <= 100 else summary.strip())
    return summaries

def generate_insight(clause_text, label=None):
    text = clause_text.lower()
    if label:
        if label == "Confidentiality": return "Ensures sensitive or proprietary information is not disclosed."
        elif label == "Termination for Convenience": return "Allows either party to terminate the agreement without specific cause."
        elif label == "Dispute Resolution": return "Describes how legal conflicts will be resolved, such as arbitration or court."
        elif label == "Cap on Liability": return "Limits the amount one party must pay for damages or losses."
        elif label == "IP Ownership": return "Specifies who holds the rights to intellectual property developed or used."
        elif label == "Anti-Assignment": return "Prevents parties from transferring their rights or obligations to others."
        elif label == "Governing Law": return "Declares which country's/state's laws apply to the contract."
        elif label == "Post-Termination": return "Describes obligations that continue even after the contract ends."
        elif label == "General legal clause": return "Standard clause included for legal completeness."
    if "terminate" in text: return "Clause allows termination of agreement."
    elif "jurisdiction" in text or "governed by" in text: return "Specifies the legal jurisdiction."
    elif "liability" in text: return "Limits financial liability."
    elif "confidential" in text: return "Protects sensitive information."
    elif "intellectual property" in text or "ownership" in text: return "Defines intellectual property rights."
    else: return "General legal clause."

def predict_clauses_batch(clauses):
    tokens = bert_tokenizer(clauses, return_tensors="pt", padding=True, truncation=True, max_length=512)
    with torch.no_grad():
        outputs = bert_model(**tokens)
        probs = torch.nn.functional.softmax(outputs.logits, dim=1)
        pred_ids = torch.argmax(probs, dim=1)
        confidences = probs.max(dim=1).values
    labels = [id2label[i.item()] for i in pred_ids]
    return labels, confidences.tolist()

def split_into_clauses(text):
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\xa0', ' ', text)
    pattern = r'(?:^|\n)(?=[0-9]+\.\s|[\u2022\-]\s|[A-Z][a-z])'
    clauses = re.split(pattern, text)
    return [c.strip() for c in clauses if len(c.strip()) > 40]

def extract_text(file):
    if file.type == "application/pdf": return extract_text_from_pdf(file)
    elif file.type == "text/plain": return file.read().decode("utf-8")
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file.type == "text/csv":
        df = pd.read_csv(file)
        return "\n".join(df.iloc[:,0].astype(str))
    elif file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        df = pd.read_excel(file)
        return "\n".join(df.iloc[:,0].astype(str))
    return ""

def extract_text_from_pdf(uploaded_pdf):
    doc = fitz.open(stream=uploaded_pdf.read(), filetype="pdf")
    return "\n".join(page.get_text() for page in doc)

def generate_file_download(dataframe, file_format):
    df_str = dataframe.astype(str)
    if file_format == "PDF":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        for _, row in df_str.iterrows():
            text = f"Clause Type: {row['Clause Type']}\nConfidence: {row['Confidence (%)']}%\nDeadline: {row['Deadline']}\nRisks: {row['Risks']}\nInsight: {row['Insight']}\nSummary: {row['Summary']}\n\n{row['Clause']}\n"
            pdf.multi_cell(0, 10, text.encode('utf-8').decode('latin-1'), border=1)
            pdf.ln()
        pdf_output = pdf.output(dest='S').encode('utf-8', 'ignore')
        return pdf_output, "application/pdf", "predicted_clauses.pdf"
    elif file_format == "TXT":
        text_data = df_str.to_string(index=False)
        return text_data.encode('utf-8'), "text/plain", "predicted_clauses.txt"
    elif file_format == "CSV":
        return df_str.to_csv(index=False).encode('utf-8'), "text/csv", "predicted_clauses.csv"
    elif file_format == "DOCX":
        doc = docx.Document()
        for _, row in df_str.iterrows():
            doc.add_paragraph(f"Clause Type: {row['Clause Type']}")
            doc.add_paragraph(f"Confidence: {row['Confidence (%)']}%")
            doc.add_paragraph(f"Deadline: {row['Deadline']}")
            doc.add_paragraph(f"Risks: {row['Risks']}")
            doc.add_paragraph(f"Insight: {row['Insight']}")
            doc.add_paragraph(f"Summary: {row['Summary']}")
            doc.add_paragraph(row['Clause'])
            doc.add_paragraph("---------------------------")
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "predicted_clauses.docx"
    elif file_format == "XLSX":
        buffer = BytesIO()
        df_str.to_excel(buffer, index=False)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "predicted_clauses.xlsx"

if selected == "Upload":
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st_lottie(lottie_contract, height=250, key="contract")
    uploaded_file = st.file_uploader("📎Upload Contract (PDF, DOCX, XLSX, TXT, CSV)", type=["pdf", "txt", "docx", "csv", "xlsx"])
    if uploaded_file:
        raw_text = extract_text(uploaded_file)
        st.text_area("Extracted Text", raw_text[:3000], height=200)
        clauses = split_into_clauses(raw_text)
        st.success(f"✓ Found {len(clauses)} potential clauses.")
        for i, c in enumerate(clauses[:5]):
            st.write(f"Clause {i+1}: {c[:100]}...")
        st.session_state.clauses = clauses
    st.markdown('</div>', unsafe_allow_html=True)

elif selected == "Analysis":
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("⌕ Run Clause Analysis") ############################################################################################3
    if "clauses" in st.session_state and st.session_state.clauses:
        if st.session_state.get("clause_data") is not None:
            st.success("✓ Analysis has already been run for this document.")
            st.info("You can view the results in the 'Insights' tab or clear the analysis from the sidebar to re-run.")
        else:
            if st.button("⌕ Analyze Clauses"): ########################################################################################
                with st.spinner("Analyzing clauses using Legal-BERT and PEGASUS-XSUM..."):
                    labels, confidences = predict_clauses_batch(st.session_state.clauses)
                    summaries = generate_summaries_batch(st.session_state.clauses)
                    data = []
                    for clause, label, conf, summary in zip(st.session_state.clauses, labels, confidences, summaries):
                        deadline, risks = extract_deadlines_and_risks(clause)
                        insight = clause_insights.get(label, "General legal clause.")
                        data.append({
                            "Clause": clause,
                            "Clause Type": label,
                            "Confidence (%)": round(conf * 100, 2),
                            "Deadline": deadline,
                            "Risks": risks,
                            "Insight": insight,
                            "Summary": summary
                        })
                    result_df = pd.DataFrame(data)
                    st.session_state.clause_data = result_df
                    persist_analysis_to_disk(result_df)
                    st.success("✓ Analysis completed.")
                    st.rerun()
    else:
        st.warning("⚠︎ No document uploaded. Go to 'Upload' to begin.")
    st.markdown('</div>', unsafe_allow_html=True)

elif selected == "Insights":
    st.markdown('<div class="card">', unsafe_allow_html=True)
    if "clause_data" not in st.session_state:
        restored = restore_analysis_from_cookie()
        if restored is not None:
            st.session_state.clause_data = restored
    if "clause_data" in st.session_state:
        result_df = st.session_state.clause_data
        st.markdown("### 📊 Overall Insights")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Clauses", len(result_df))
        with col2:
            st.metric("Unique Clause Types", len(result_df["Clause Type"].unique()))
        with col3:
            st.metric("High-Risk Clauses", (result_df["Risks"] != "-").sum())
        with col4:
            st.metric("Clauses with Deadlines", (result_df["Deadline"] != "-").sum())

        st.markdown("---")
        st.markdown("### 📊 Full Clause Insights")
        st.dataframe(result_df, use_container_width=True)

        st.markdown("### ⚠︎ High Risk Clauses") #################################################################################
        st.dataframe(result_df[result_df["Risks"] != "-"][["Clause", "Risks"]], use_container_width=True)

        st.markdown("### ⏱ Clauses with Deadlines")
        st.dataframe(result_df[result_df["Deadline"] != "-"][["Clause", "Deadline"]], use_container_width=True)

        st.markdown("### 📝 Clause Summaries")
        st.dataframe(result_df[["Clause", "Summary"]], use_container_width=True)

        if result_df["Summary"].eq(result_df["Clause"]).all():
            st.warning("⚠︎ All summaries are identical to clauses. Try uploading longer clauses or more complex contracts.")
    else:
        st.warning("⚠︎ No analyzed data found. Go to 'Analysis' tab.") ######################################################
    st.markdown('</div>', unsafe_allow_html=True)

elif selected == "Download":
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("📥 Download Report")
    if "clause_data" in st.session_state and st.session_state.clause_data is not None and not st.session_state.clause_data.empty:
        result_df = st.session_state.clause_data
        file_format = st.selectbox("Select Format", ["PDF", "DOCX", "XLSX", "TXT", "CSV"])
        file_data, mime, filename = generate_file_download(result_df, file_format)
        st.download_button(
            f"⬇️ Download Clause Report ({file_format})",
            file_data,
            filename,
            mime
        )
    else:
        st.warning("⚠︎ Nothing to download. Run analysis first.")
    st.markdown('</div>', unsafe_allow_html=True)

# --- Sidebar: Clause Repository ---
st.sidebar.title("📚 Clause Repository")

# 🔹 Add Clear Saved Analysis button at the top
if st.sidebar.button("🗑️ Clear Saved Analysis"):
    clear_persisted_analysis()
    st.session_state.pop("clause_data", None)
    st.success("Analysis data cleared.")

if "clause_data" in st.session_state:
    clause_df = st.session_state.clause_data
    clause_types = sorted(clause_df["Clause Type"].unique())

    # --- Filters ---
    selected_type = st.sidebar.selectbox("⌕ Filter by Clause Type", ["All"] + clause_types)
    search_query = st.sidebar.text_input("⌕ Search Clause or Summary")

    # --- Apply Filters ---
    filtered_df = clause_df.copy()
    if selected_type != "All":
        filtered_df = filtered_df[filtered_df["Clause Type"] == selected_type]
    if search_query:
        filtered_df = filtered_df[
            filtered_df["Clause"].str.contains(search_query, case=False, na=False) |
            filtered_df["Summary"].str.contains(search_query, case=False, na=False)
        ]

    st.sidebar.markdown(f"**🗐 Total Filtered Clauses: {len(filtered_df)}**")

    # --- Display Clauses with Full Info ---
    for i, row in filtered_df.iterrows():
        with st.sidebar.expander(f"🗐 Clause {i+1}: {row['Clause Type']}"):
            st.markdown(f"**Clause:** {row['Clause']}")
            st.markdown(f"**Summary:** {row['Summary']}")
            st.markdown(f"**Insight:** {row['Insight']}")
            st.markdown(f"**Risks:** {row['Risks'] if row['Risks'] != '-' else 'No specific risk detected'}**")

            # ✅ More meaningful deadline display
            if row['Deadline'] and row['Deadline'] != "-":
                st.markdown(f"**Deadline:** {row['Deadline']} 📅")
            else:
                st.markdown("**Deadline:** ❌ Not specified")

            st.markdown(f"**Confidence:** {row['Confidence (%)']}%")
else:
    st.sidebar.info("🗐 Upload and analyze a document to view repository.")



